from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


OUTPUT_DIR = Path("/Users/pietrolimoni/Desktop/Jump-Analysis/outputs/manual-presentation-part2/final")
PDF_PATH = OUTPUT_DIR / "part2_sensors_data_collection_sync.pdf"
DOCX_PATH = OUTPUT_DIR / "part2_speech_pietro.docx"

PAGE_SIZE = landscape((13.333 * 72, 7.5 * 72))
PAGE_W, PAGE_H = PAGE_SIZE

TOP_BAND_H = 80
BOTTOM_BAND_H = 52
MARGIN_X = 42
CONTENT_TOP = PAGE_H - TOP_BAND_H - 26
CONTENT_BOTTOM = BOTTOM_BAND_H + 24

LIGHT_BAND = colors.HexColor("#efefef")
TITLE_WHITE = colors.white
TEXT_DARK = colors.HexColor("#222222")
TEXT_MUTED = colors.HexColor("#555555")
ACCENT = colors.HexColor("#8b0000")
BOX_FILL = colors.HexColor("#f7f7f7")
BOX_STROKE = colors.HexColor("#cfcfcf")


def draw_template_frame(pdf: canvas.Canvas, title: str, slide_number: int, total_slides: int) -> None:
    pdf.setFillColor(LIGHT_BAND)
    pdf.rect(0, PAGE_H - TOP_BAND_H, PAGE_W, TOP_BAND_H, fill=1, stroke=0)
    pdf.rect(0, 0, PAGE_W, BOTTOM_BAND_H, fill=1, stroke=0)

    pdf.setFillColor(TITLE_WHITE)
    pdf.setFont("Helvetica-BoldOblique", 27)
    pdf.drawString(16, PAGE_H - 38, title)

    pdf.setFillColor(TEXT_DARK)
    pdf.setFont("Helvetica", 18)
    pdf.drawRightString(PAGE_W - 20, 17, f"{slide_number}/{total_slides}")

    pdf.setFillColor(TEXT_MUTED)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(12, 16, "Politecnico di Milano")


def draw_title_slide(pdf: canvas.Canvas, total_slides: int) -> None:
    pdf.setFillColor(LIGHT_BAND)
    pdf.rect(0, PAGE_H - TOP_BAND_H, PAGE_W, TOP_BAND_H, fill=1, stroke=0)
    pdf.rect(0, 0, PAGE_W, BOTTOM_BAND_H + 14, fill=1, stroke=0)
    pdf.rect(16, PAGE_H - 140, 360, 105, fill=1, stroke=0)

    pdf.setFillColor(TEXT_DARK)
    pdf.setFont("Helvetica-Bold", 28)
    pdf.drawCentredString(PAGE_W / 2, PAGE_H / 2 + 50, "Sensors, Data Collection,")
    pdf.drawCentredString(PAGE_W / 2, PAGE_H / 2 + 14, "and Synchronization")

    pdf.setFillColor(TEXT_MUTED)
    pdf.setFont("Helvetica", 17)
    pdf.drawCentredString(PAGE_W / 2, PAGE_H / 2 - 28, "Section 2 of the presentation")
    pdf.drawCentredString(PAGE_W / 2, PAGE_H / 2 - 52, "Pietro Limoni")

    pdf.setFillColor(TEXT_MUTED)
    pdf.setFont("Helvetica", 11)
    pdf.drawCentredString(
        PAGE_W / 2,
        18,
        "Scuola di Ingegneria Industriale e dell'Informazione  |  Thesis Presentation",
    )


def draw_bullets(pdf: canvas.Canvas, x: float, y: float, bullets: list[str], width: float, font_size: int = 20) -> None:
    cursor_y = y
    for bullet in bullets:
        wrapped_lines = wrap_text(bullet, width - 26, font_size, "Helvetica")
        pdf.setFillColor(ACCENT)
        pdf.circle(x + 5, cursor_y + 4, 3, fill=1, stroke=0)
        pdf.setFillColor(TEXT_DARK)
        for line_index, line in enumerate(wrapped_lines):
            pdf.setFont("Helvetica", font_size)
            pdf.drawString(x + 18, cursor_y - (line_index * (font_size + 4)), line)
        cursor_y -= max(font_size + 16, len(wrapped_lines) * (font_size + 6) + 8)


def wrap_text(text: str, width: float, font_size: int, font_name: str) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if stringWidth(candidate, font_name, font_size) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_info_box(pdf: canvas.Canvas, x: float, y: float, w: float, h: float, title: str, body: str) -> None:
    pdf.setFillColor(BOX_FILL)
    pdf.setStrokeColor(BOX_STROKE)
    pdf.roundRect(x, y, w, h, 12, fill=1, stroke=1)
    pdf.setFillColor(ACCENT)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(x + 14, y + h - 24, title)
    pdf.setFillColor(TEXT_DARK)
    pdf.setFont("Helvetica", 14)
    lines = wrap_text(body, w - 28, 14, "Helvetica")
    text_y = y + h - 46
    for line in lines[:6]:
        pdf.drawString(x + 14, text_y, line)
        text_y -= 18


def draw_pipeline_slide(pdf: canvas.Canvas) -> None:
    draw_template_frame(pdf, "Data Collection Pipeline", 3, 5)

    steps = [
        "Athlete height",
        "Floor setup",
        "Box setup",
        "Drop detection",
        "Protocol validation",
        "Save trial",
    ]
    x0 = 55
    y = PAGE_H / 2 + 10
    box_w = 130
    gap = 18
    for index, step in enumerate(steps):
        x = x0 + index * (box_w + gap)
        pdf.setFillColor(BOX_FILL)
        pdf.setStrokeColor(BOX_STROKE)
        pdf.roundRect(x, y, box_w, 64, 10, fill=1, stroke=1)
        pdf.setFillColor(TEXT_DARK)
        pdf.setFont("Helvetica-Bold", 16)
        lines = wrap_text(step, box_w - 16, 16, "Helvetica-Bold")
        line_y = y + 39
        for line in lines:
            pdf.drawCentredString(x + box_w / 2, line_y, line)
            line_y -= 18
        if index < len(steps) - 1:
            arrow_x = x + box_w
            pdf.setStrokeColor(ACCENT)
            pdf.setLineWidth(2)
            pdf.line(arrow_x + 6, y + 32, arrow_x + gap - 6, y + 32)
            pdf.line(arrow_x + gap - 12, y + 36, arrow_x + gap - 6, y + 32)
            pdf.line(arrow_x + gap - 12, y + 28, arrow_x + gap - 6, y + 32)

    bullets = [
        "The pipeline follows the same logic as the main jump-analysis workflow.",
        "Recording starts automatically only when the real drop is detected.",
        "If setup or protocol checks fail, the trial is rejected and not saved.",
    ]
    draw_bullets(pdf, 90, 220, bullets, PAGE_W - 180, 18)


def draw_saved_data_slide(pdf: canvas.Canvas) -> None:
    draw_template_frame(pdf, "What We Save for Each Valid Trial", 4, 5)

    draw_info_box(
        pdf,
        52,
        290,
        255,
        210,
        "movement_timeseries.csv",
        "Frame-by-frame pose trajectory, video-derived knee orientation proxies, aligned left/right IMU pitch-roll-yaw, and body-scale values over time.",
    )
    draw_info_box(
        pdf,
        332,
        290,
        255,
        210,
        "front_2d_features.csv",
        "The 37 frontal features used by the vision pipeline, extracted only from valid trials.",
    )
    draw_info_box(
        pdf,
        612,
        290,
        255,
        210,
        "trial_metadata.json",
        "Setup calibration, protocol checks, participant settings, and sensor acquisition status.",
    )
    draw_info_box(
        pdf,
        892,
        290,
        255,
        210,
        "Raw IMU CSV files",
        "When live sensors are used, left and right raw BWT901CL recordings are also stored for later analysis.",
    )

    bullets = [
        "Each trial folder combines vision data and sensor data in one structured package.",
        "This keeps collection separate from final model inference.",
    ]
    draw_bullets(pdf, 90, 200, bullets, PAGE_W - 180, 18)


def draw_sync_slide(pdf: canvas.Canvas) -> None:
    draw_template_frame(pdf, "Synchronization and Why It Matters", 5, 5)

    pdf.setFillColor(TEXT_DARK)
    pdf.setFont("Helvetica-Bold", 21)
    pdf.drawString(74, 560, "How alignment works")
    pdf.drawString(690, 560, "Why this matters")

    left_bullets = [
        "Video and IMUs can have different sampling rates.",
        "The first video frame is used as the jump start reference.",
        "IMU timestamps are converted to time-from-start.",
        "Sensor values are interpolated onto video timestamps.",
        "Each video frame gets matched sensor measurements.",
    ]
    right_bullets = [
        "The dataset is cleaner because invalid trials are filtered out.",
        "The IMUs provide ground truth for future supervised learning.",
        "Later models can correct or replace video-only knee orientation estimates.",
        "So this part builds reliable training data, not just recordings.",
    ]
    draw_bullets(pdf, 74, 520, left_bullets, 500, 17)
    draw_bullets(pdf, 690, 520, right_bullets, 500, 17)

    pdf.setStrokeColor(BOX_STROKE)
    pdf.setLineWidth(1)
    pdf.line(PAGE_W / 2, 135, PAGE_W / 2, 545)


def build_pdf() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(PDF_PATH), pagesize=PAGE_SIZE)

    draw_title_slide(pdf, 5)
    pdf.showPage()

    draw_template_frame(pdf, "Why We Added Sensors", 2, 5)
    draw_bullets(
        pdf,
        75,
        520,
        [
            "A frontal 2D camera is useful, but it cannot give reliable ground truth for knee orientation.",
            "So we added two BWT901CL IMU sensors, one near each knee.",
            "The goal is to collect synchronized video and IMU data during the drop jump.",
            "These sensor signals become the reference for future supervised learning.",
        ],
        620,
        19,
    )
    draw_info_box(
        pdf,
        760,
        300,
        400,
        170,
        "Key idea",
        "Video estimates movement. IMUs provide a more reliable reference for knee pitch, roll, and yaw.",
    )
    draw_info_box(
        pdf,
        760,
        130,
        400,
        120,
        "Sensor setup",
        "Two BWT901CL sensors placed near the knees during the drop-jump recording.",
    )
    pdf.showPage()

    draw_pipeline_slide(pdf)
    pdf.showPage()

    draw_saved_data_slide(pdf)
    pdf.showPage()

    draw_sync_slide(pdf)
    pdf.showPage()

    pdf.save()


def build_docx() -> None:
    speech = [
        "My part is about sensors, data collection, and synchronization.",
        "The main idea is this: the camera alone is useful, but it cannot give us reliable ground-truth information about knee orientation. From a frontal 2D video, we can estimate movement, but we cannot measure pitch, roll, and yaw accurately enough to use them as a reference.",
        "So, to improve this, we added two IMU sensors, one on each knee. In our project, we use BWT901CL sensors. Their role is to record orientation during the drop jump, so we can pair video information with sensor information.",
        "The data collection pipeline follows the same logic as the main jump analysis workflow. First, the athlete enters their height. Then the system performs a setup phase: first on the floor, and then on the box. This step is important because it helps the system estimate scale and verify that the camera position is acceptable.",
        "After setup, the athlete stands on the box and waits. Recording does not start manually. Instead, it starts automatically when the system detects the actual drop. This is useful because we want to capture the right movement window and avoid unnecessary frames.",
        "Another important point is data quality. We do not save every attempt. A trial is saved only if the setup is correct and if the jump follows the expected protocol. For example, the athlete must really start from the box, land correctly, and perform the rebound jump. If these conditions are not respected, the trial is rejected. This way, we collect cleaner data.",
        "For each valid trial, we save several things. We save the frame-by-frame pose trajectory from the video, the extracted 2D features, the video-based knee orientation estimates, and the left and right IMU signals. We also save metadata, such as calibration information and protocol checks.",
        "The key technical part is synchronization. Video and IMU sensors usually do not have the same sampling rate, so we align them using time. We take the beginning of the jump as a reference, convert the IMU data to the same time scale, and then interpolate the sensor values onto the video timestamps. In this way, each video frame can be matched with the corresponding sensor measurements.",
        "So the main contribution of this part is not just adding sensors, but building a clean and synchronized dataset. This dataset can then be used later to train supervised models, using the IMU data as ground truth.",
    ]

    document = Document()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Part 2 Speech - Sensors, Data Collection, and Synchronization")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Simplified spoken version for presenting naturally")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)

    for paragraph_text in speech:
        paragraph = document.add_paragraph(paragraph_text)
        paragraph.style = document.styles["Normal"]
        paragraph.paragraph_format.space_after = Pt(8)

    document.save(DOCX_PATH)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
